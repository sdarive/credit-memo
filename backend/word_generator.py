"""
Word Document Generator for Credit Memos
Generates professionally formatted .docx files with auto-populated fields
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from datetime import datetime
import os


class CreditMemoWordGenerator:
    """
    Generates professional credit memo Word documents following
    Midwest Regional Bank template standards
    """

    def __init__(self):
        self.document = Document()
        self._setup_styles()

    def _setup_styles(self):
        """Configure document styles"""
        # Set default font
        style = self.document.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        # Margins (1 inch all sides)
        sections = self.document.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

    def generate_credit_memo(self, extracted_data, ratios, memo_narrative,
                            borrower_info, output_path):
        """
        Generate complete credit memo Word document

        Args:
            extracted_data: Dict of financial data from ADE
            ratios: Dict of calculated financial ratios
            memo_narrative: Generated credit memo text
            borrower_info: Dict with borrower_name, industry, loan_amount, etc.
            output_path: Path to save .docx file
        """

        # Add header
        self._add_header()

        # Add confidential marking
        self._add_confidential_marking()

        # Add document title
        self._add_title("CREDIT MEMORANDUM")

        # Add loan information section
        self._add_loan_info_section(borrower_info, extracted_data)

        # Add executive summary
        self._add_section_heading("EXECUTIVE SUMMARY & RECOMMENDATION")
        self._add_executive_summary(borrower_info, ratios)

        # Add 5 C's analysis
        self._add_five_cs_analysis(extracted_data, ratios, borrower_info)

        # Add financial performance table
        self._add_section_heading("HISTORICAL FINANCIAL PERFORMANCE")
        self._add_financial_table(extracted_data)

        # Add financial ratios table
        self._add_section_heading("FINANCIAL RATIOS ANALYSIS")
        self._add_ratios_table(ratios)

        # Add risk assessment
        self._add_section_heading("RISK ASSESSMENT")
        self._add_risk_assessment(ratios)

        # Add strengths and concerns
        self._add_strengths_concerns(ratios)

        # Add recommendation
        self._add_section_heading("RECOMMENDATION")
        self._add_recommendation(borrower_info, ratios)

        # Add signature block
        self._add_signature_block()

        # Add footer
        self._add_footer()

        # Save document
        self.document.save(output_path)

        return output_path

    def _add_header(self):
        """Add document header with bank name and branding"""
        section = self.document.sections[0]
        header = section.header

        # Bank name
        bank_para = header.paragraphs[0]
        bank_para.text = "MIDWEST REGIONAL BANK"
        bank_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        bank_run = bank_para.runs[0]
        bank_run.font.size = Pt(14)
        bank_run.font.bold = True
        bank_run.font.color.rgb = RGBColor(0, 103, 71)  # Banking green (#006747)

        # Division name
        division = header.add_paragraph()
        division.text = "Business & Commercial Finance"
        division.alignment = WD_ALIGN_PARAGRAPH.CENTER
        division_run = division.runs[0]
        division_run.font.size = Pt(11)
        division_run.font.italic = False
        division_run.font.color.rgb = RGBColor(51, 51, 51)  # Dark gray

        # FDIC tagline
        fdic = header.add_paragraph()
        fdic.text = "Member FDIC"
        fdic.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fdic_run = fdic.runs[0]
        fdic_run.font.size = Pt(9)
        fdic_run.font.italic = True
        fdic_run.font.color.rgb = RGBColor(102, 102, 102)  # Medium gray

        # Add separator line
        header.add_paragraph("_" * 80).alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_confidential_marking(self):
        """Add confidential marking"""
        conf_para = self.document.add_paragraph()
        conf_para.text = "CONFIDENTIAL - INTERNAL USE ONLY"
        conf_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        conf_run = conf_para.runs[0]
        conf_run.font.bold = True
        conf_run.font.size = Pt(10)
        conf_run.font.color.rgb = RGBColor(255, 0, 0)  # Red

        self.document.add_paragraph()  # Spacing

    def _add_title(self, title_text):
        """Add document title"""
        title = self.document.add_paragraph()
        title.text = title_text
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.size = Pt(16)
        title_run.font.bold = True

        self.document.add_paragraph()  # Spacing

    def _add_loan_info_section(self, borrower_info, extracted_data):
        """Add loan information table"""
        table = self.document.add_table(rows=9, cols=2)
        table.style = 'Light Grid Accent 1'

        # Set column widths
        table.columns[0].width = Inches(2.5)
        table.columns[1].width = Inches(4.0)

        # Populate table
        info_items = [
            ("Borrower:", borrower_info.get('borrower_name', 'N/A')),
            ("Industry:", borrower_info.get('industry', 'N/A')),
            ("Loan Type:", borrower_info.get('loan_type', 'Commercial Term Loan')),
            ("Requested Amount:", self._format_currency(borrower_info.get('loan_amount', 0))),
            ("Purpose:", borrower_info.get('purpose', 'Working capital and business expansion')),
            ("Date:", datetime.now().strftime("%B %d, %Y")),
            ("Loan Officer:", borrower_info.get('loan_officer', '[Loan Officer Name]')),
            ("Credit Analyst:", borrower_info.get('credit_analyst', '[Credit Analyst Name]')),
            ("Risk Rating:", self._determine_risk_rating(extracted_data))
        ]

        for i, (label, value) in enumerate(info_items):
            row = table.rows[i]
            row.cells[0].text = label
            row.cells[0].paragraphs[0].runs[0].font.bold = True
            row.cells[1].text = str(value)

        self.document.add_paragraph()  # Spacing

    def _add_section_heading(self, heading_text):
        """Add formatted section heading"""
        self.document.add_paragraph()  # Spacing before

        heading = self.document.add_heading(heading_text, level=1)
        heading_run = heading.runs[0]
        heading_run.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
        heading_run.font.size = Pt(14)

    def _add_executive_summary(self, borrower_info, ratios):
        """Add executive summary with recommendation"""
        # Recommendation badge
        recommendation = self._determine_recommendation(ratios)
        rec_para = self.document.add_paragraph()
        rec_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rec_run = rec_para.add_run(f"RECOMMENDATION: {recommendation}")
        rec_run.font.bold = True
        rec_run.font.size = Pt(12)

        # Color code recommendation
        if "APPROVED" in recommendation:
            rec_run.font.color.rgb = RGBColor(0, 128, 0)  # Green
        elif "DECLINED" in recommendation:
            rec_run.font.color.rgb = RGBColor(255, 0, 0)  # Red
        else:
            rec_run.font.color.rgb = RGBColor(255, 165, 0)  # Orange

        self.document.add_paragraph()  # Spacing

        # Summary narrative
        summary = self.document.add_paragraph()
        summary.add_run(
            f"{borrower_info.get('borrower_name', 'The borrower')} requests "
            f"{self._format_currency(borrower_info.get('loan_amount', 0))} for "
            f"{borrower_info.get('purpose', 'business purposes')}. "
        )

        # Add DSCR highlight
        dscr = ratios.get('dscr', 0)
        if isinstance(dscr, dict):
            dscr = dscr.get('value', 0)

        if dscr and isinstance(dscr, (int, float)):
            summary.add_run(
                f"The company demonstrates a Debt Service Coverage Ratio of {dscr:.2f}x, "
            )

            if dscr >= 1.25:
                summary.add_run("indicating strong repayment capacity. ")
            elif dscr >= 1.0:
                summary.add_run("indicating adequate but tight repayment capacity. ")
            else:
                summary.add_run("raising concerns about repayment capacity. ")

        # Primary repayment source
        repayment = self.document.add_paragraph()
        repayment.add_run("Primary Repayment Source: ").bold = True
        repayment.add_run("Operating cash flow from business operations")

        # Secondary repayment source
        secondary = self.document.add_paragraph()
        secondary.add_run("Secondary Repayment Source: ").bold = True
        secondary.add_run("Liquidation of business assets and personal guarantees")

    def _add_five_cs_analysis(self, extracted_data, ratios, borrower_info):
        """Add 5 C's of Credit analysis"""

        # 1. CHARACTER
        self._add_section_heading("THE 5 C'S ANALYSIS")

        char_heading = self.document.add_heading("1. CHARACTER", level=2)
        char_heading.runs[0].font.size = Pt(12)

        char_para = self.document.add_paragraph()
        char_para.add_run("Credit History: ").bold = True
        char_para.add_run(
            f"{borrower_info.get('borrower_name', 'The borrower')} has "
            f"maintained banking relationships with strong payment history. "
            f"No prior bankruptcies, liens, or judgments identified."
        )

        # 2. CAPACITY
        cap_heading = self.document.add_heading("2. CAPACITY", level=2)
        cap_heading.runs[0].font.size = Pt(12)

        cap_para = self.document.add_paragraph()
        cap_para.add_run("Cash Flow Analysis: ").bold = True

        revenue = extracted_data.get('revenue', 0)
        net_income = extracted_data.get('net_income', 0)
        ebitda = extracted_data.get('ebitda', 0)

        cap_para.add_run(
            f"The company generated {self._format_currency(revenue)} in revenue "
            f"with net income of {self._format_currency(net_income)}. "
            f"EBITDA of {self._format_currency(ebitda)} demonstrates "
        )

        if ebitda and revenue and ebitda > revenue * 0.15:
            cap_para.add_run("strong operating profitability.")
        elif ebitda and revenue and ebitda > revenue * 0.08:
            cap_para.add_run("adequate operating profitability.")
        else:
            cap_para.add_run("concerning profitability levels requiring attention.")

        # 3. CAPITAL
        cap_heading = self.document.add_heading("3. CAPITAL", level=2)
        cap_heading.runs[0].font.size = Pt(12)

        capital_para = self.document.add_paragraph()
        capital_para.add_run("Equity Position: ").bold = True

        assets = extracted_data.get('total_assets', 0)
        liabilities = extracted_data.get('total_liabilities', 0)
        equity = assets - liabilities if assets and liabilities else 0

        capital_para.add_run(
            f"Total assets of {self._format_currency(assets)} with "
            f"liabilities of {self._format_currency(liabilities)} result in "
            f"equity of {self._format_currency(equity)}. "
        )

        # 4. COLLATERAL
        coll_heading = self.document.add_heading("4. COLLATERAL", level=2)
        coll_heading.runs[0].font.size = Pt(12)

        coll_para = self.document.add_paragraph()
        coll_para.add_run("Collateral Structure: ").bold = True
        coll_para.add_run(
            "Loan secured by [collateral description]. "
            "UCC-1 blanket lien on all business assets. "
            "Personal guarantees from principal owners."
        )

        # 5. CONDITIONS
        cond_heading = self.document.add_heading("5. CONDITIONS", level=2)
        cond_heading.runs[0].font.size = Pt(12)

        cond_para = self.document.add_paragraph()
        cond_para.add_run("Loan Terms: ").bold = True
        cond_para.add_run(
            f"Proposed loan amount: {self._format_currency(borrower_info.get('loan_amount', 0))}. "
            f"Interest rate: Prime + 2.50% (variable). "
            f"Term: 5-7 years with amortization schedule. "
            f"Personal guarantees required from all owners ≥20%."
        )

    def _add_financial_table(self, extracted_data):
        """Add historical financial performance table"""
        table = self.document.add_table(rows=6, cols=2)
        table.style = 'Light Shading Accent 1'

        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "Financial Metric"
        header_cells[1].text = "Most Recent Period"

        for cell in header_cells:
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Data rows
        financial_items = [
            ("Revenue", extracted_data.get('revenue', 0)),
            ("Net Income", extracted_data.get('net_income', 0)),
            ("EBITDA", extracted_data.get('ebitda', 0)),
            ("Total Assets", extracted_data.get('total_assets', 0)),
            ("Total Liabilities", extracted_data.get('total_liabilities', 0))
        ]

        for i, (label, value) in enumerate(financial_items, start=1):
            row = table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = self._format_currency(value)
            row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        self.document.add_paragraph()  # Spacing

    def _add_ratios_table(self, ratios):
        """Add financial ratios analysis table with color coding"""
        table = self.document.add_table(rows=10, cols=4)
        table.style = 'Light Grid Accent 1'

        # Header row
        headers = ["Ratio", "Value", "Threshold", "Status"]
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Ratio definitions
        ratio_info = [
            ("Debt Service Coverage (DSCR)", "dscr", "≥1.25x"),
            ("Total Debt to EBITDA", "debt_to_ebitda", "≤2.0x"),
            ("Current Ratio", "current_ratio", "≥2.0x"),
            ("Quick Ratio", "quick_ratio", "≥1.0x"),
            ("Net Income Margin", "net_income_margin", "≥10%"),
            ("Interest Coverage", "interest_coverage", "≥3.0x"),
            ("Leverage Ratio", "leverage_ratio", "≤0.3"),
            ("Working Capital", "working_capital", ">$0"),
            ("Days Sales Outstanding", "dso", "≤45 days")
        ]

        for i, (ratio_name, ratio_key, threshold) in enumerate(ratio_info, start=1):
            row = table.rows[i]

            # Ratio name
            row.cells[0].text = ratio_name

            # Get ratio value (handle both dict and flat formats)
            ratio_data = ratios.get(ratio_key, None)
            if isinstance(ratio_data, dict):
                value = ratio_data.get('value', 'N/A')
                status = ratio_data.get('status', 'unknown')
            else:
                value = ratio_data
                status = self._get_ratio_status(ratio_key, value)

            # Value
            if ratio_key == 'net_income_margin':
                value_text = f"{value}%" if isinstance(value, (int, float)) else str(value)
            elif ratio_key == 'working_capital':
                value_text = self._format_currency(value) if isinstance(value, (int, float)) else str(value)
            elif ratio_key == 'dso':
                value_text = f"{value} days" if isinstance(value, (int, float)) else str(value)
            else:
                value_text = f"{value:.2f}x" if isinstance(value, (int, float)) else str(value)

            row.cells[1].text = value_text
            row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Threshold
            row.cells[2].text = threshold
            row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Status with color
            status_cell = row.cells[3]
            status_para = status_cell.paragraphs[0]
            status_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            status_run = status_para.add_run(status.upper())
            status_run.font.bold = True

            if status in ['healthy', 'good']:
                status_run.font.color.rgb = RGBColor(0, 128, 0)  # Green
            elif status in ['watch', 'fair']:
                status_run.font.color.rgb = RGBColor(255, 165, 0)  # Orange
            elif status in ['concern', 'poor']:
                status_run.font.color.rgb = RGBColor(255, 0, 0)  # Red

        self.document.add_paragraph()  # Spacing

    def _add_risk_assessment(self, ratios):
        """Add risk assessment narrative"""
        risk_para = self.document.add_paragraph()

        # Count healthy vs concerning ratios
        healthy_count = 0
        concern_count = 0

        for ratio_value in ratios.values():
            if isinstance(ratio_value, dict):
                status = ratio_value.get('status', '')
            else:
                status = ''

            if status in ['healthy', 'good']:
                healthy_count += 1
            elif status in ['concern', 'poor']:
                concern_count += 1

        if concern_count == 0 and healthy_count >= 7:
            risk_level = "LOW"
            risk_desc = "demonstrates strong financial performance across all key metrics"
        elif concern_count <= 2:
            risk_level = "MODERATE"
            risk_desc = "shows acceptable performance with some areas requiring monitoring"
        else:
            risk_level = "HIGH"
            risk_desc = "exhibits concerning financial trends requiring enhanced oversight"

        risk_para.add_run(f"Overall Risk Assessment: ").bold = True
        risk_run = risk_para.add_run(risk_level)
        risk_run.bold = True

        if risk_level == "LOW":
            risk_run.font.color.rgb = RGBColor(0, 128, 0)
        elif risk_level == "MODERATE":
            risk_run.font.color.rgb = RGBColor(255, 165, 0)
        else:
            risk_run.font.color.rgb = RGBColor(255, 0, 0)

        risk_para.add_run(f". The borrower {risk_desc}.")

    def _add_strengths_concerns(self, ratios):
        """Add strengths and concerns lists"""
        # Strengths
        strengths_heading = self.document.add_heading("STRENGTHS", level=2)
        strengths_heading.runs[0].font.size = Pt(12)
        strengths_heading.runs[0].font.color.rgb = RGBColor(0, 128, 0)

        strengths = []
        interpretations = ratios.get('interpretations', {})

        for ratio_key, interpretation in interpretations.items():
            if ratio_key != 'interpretations':
                # Check if this is a positive interpretation
                if any(word in interpretation.lower() for word in ['strong', 'excellent', 'good', 'healthy', 'positive']):
                    strengths.append(interpretation)

        if not strengths:
            strengths = [
                "Established business with operating history",
                "Experienced management team",
                "Banking relationship maintained"
            ]

        for strength in strengths[:5]:  # Limit to 5
            self.document.add_paragraph(strength, style='List Bullet')

        # Concerns
        concerns_heading = self.document.add_heading("CONCERNS & MITIGATION", level=2)
        concerns_heading.runs[0].font.size = Pt(12)
        concerns_heading.runs[0].font.color.rgb = RGBColor(255, 0, 0)

        concerns = []
        for ratio_key, interpretation in interpretations.items():
            if ratio_key != 'interpretations':
                # Check if this is a negative interpretation
                if any(word in interpretation.lower() for word in ['weak', 'concern', 'low', 'risk', 'slow', 'negative']):
                    concerns.append(interpretation)

        if not concerns:
            concerns = [
                "Economic conditions may impact industry performance",
                "Competition in market requires ongoing monitoring",
                "Regular financial reporting required to track trends"
            ]

        for concern in concerns[:5]:  # Limit to 5
            self.document.add_paragraph(concern, style='List Bullet')

    def _add_recommendation(self, borrower_info, ratios):
        """Add final recommendation section"""
        rec_para = self.document.add_paragraph()

        recommendation = self._determine_recommendation(ratios)

        rec_para.add_run(f"{recommendation}: ").bold = True

        loan_amount = self._format_currency(borrower_info.get('loan_amount', 0))
        borrower_name = borrower_info.get('borrower_name', 'the borrower')

        rec_para.add_run(
            f"Loan of {loan_amount} to {borrower_name} "
        )

        if "APPROVED" in recommendation:
            rec_para.add_run(
                "based on demonstrated financial capacity, adequate collateral coverage, "
                "and acceptable risk profile. "
            )
        else:
            rec_para.add_run(
                "due to financial performance concerns requiring additional analysis "
                "and risk mitigation strategies. "
            )

        # Conditions precedent
        conditions_heading = self.document.add_heading("Conditions Precedent:", level=3)
        conditions_heading.runs[0].font.size = Pt(11)

        standard_conditions = [
            "Execution of loan agreement and security documents",
            "Personal guarantees from all owners ≥20%",
            "UCC-1 financing statements filed and perfected",
            "Proof of insurance with bank named as loss payee",
            "Receipt of current financial statements (≤90 days old)",
            "Board resolution authorizing borrowing (if applicable)"
        ]

        for condition in standard_conditions:
            self.document.add_paragraph(condition, style='List Bullet')

        # Ongoing monitoring
        monitoring_heading = self.document.add_heading("Ongoing Monitoring:", level=3)
        monitoring_heading.runs[0].font.size = Pt(11)

        monitoring_items = [
            "Annual financial statements (CPA-reviewed) due within 90 days of fiscal year-end",
            "Quarterly internal financial statements",
            "Annual covenant compliance testing",
            "Periodic site visits and management meetings",
            "Risk rating review with each financial submission"
        ]

        for item in monitoring_items:
            self.document.add_paragraph(item, style='List Bullet')

    def _add_signature_block(self):
        """Add signature block for approvals"""
        self.document.add_page_break()

        sig_heading = self.document.add_heading("APPROVAL SIGNATURES", level=2)
        sig_heading.runs[0].font.size = Pt(12)

        self.document.add_paragraph()
        self.document.add_paragraph()

        # Signature lines
        signatures = [
            "Prepared by: ________________________________   Date: ______________",
            "                    Credit Analyst",
            "",
            "Reviewed by: ________________________________   Date: ______________",
            "                    Chief Credit Officer",
            "",
            "Approved by: ________________________________   Date: ______________",
            "                    President & CEO"
        ]

        for sig_line in signatures:
            para = self.document.add_paragraph(sig_line)
            if ":" not in sig_line and sig_line.strip():
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.runs[0].font.size = Pt(10)
                para.runs[0].font.italic = True

    def _add_footer(self):
        """Add document footer with branding"""
        section = self.document.sections[0]
        footer = section.footer

        # Primary footer line
        footer_para = footer.paragraphs[0]
        footer_para.text = "Document Control: Internal Confidential | Classification: Credit Documentation"
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_para.runs[0]
        footer_run.font.size = Pt(8)
        footer_run.font.italic = True
        footer_run.font.color.rgb = RGBColor(128, 128, 128)

        # Ernie branding line
        ernie_para = footer.add_paragraph()
        ernie_para.text = "Generated by Ernie - AI Credit Assistant | Powered by LandingAI & AWS Bedrock"
        ernie_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ernie_run = ernie_para.runs[0]
        ernie_run.font.size = Pt(7)
        ernie_run.font.italic = True
        ernie_run.font.color.rgb = RGBColor(150, 150, 150)

    # Helper methods

    def _format_currency(self, value):
        """Format number as currency"""
        try:
            if value is None:
                return "N/A"
            return f"${value:,.0f}"
        except (ValueError, TypeError):
            return "N/A"

    def _determine_risk_rating(self, extracted_data):
        """Determine risk rating based on financial metrics"""
        # Simplified risk rating logic
        # In production, use more sophisticated analysis
        return "4 (Pass - Acceptable Risk)"

    def _determine_recommendation(self, ratios):
        """Determine loan recommendation based on ratios"""
        concern_count = 0
        healthy_count = 0

        interpretations = ratios.get('interpretations', {})
        for ratio_key, interpretation in interpretations.items():
            if any(word in interpretation.lower() for word in ['strong', 'excellent', 'good', 'healthy', 'positive']):
                healthy_count += 1
            elif any(word in interpretation.lower() for word in ['weak', 'concern', 'low', 'risk']):
                concern_count += 1

        if concern_count == 0 and healthy_count >= 7:
            return "APPROVED"
        elif concern_count >= 3:
            return "DECLINED"
        else:
            return "APPROVED WITH CONDITIONS"

    def _get_ratio_status(self, ratio_key, value):
        """Determine status for a ratio value"""
        if value is None or not isinstance(value, (int, float)):
            return 'unknown'

        # Define thresholds for each ratio
        thresholds = {
            'dscr': {'healthy': 1.25, 'watch': 1.0},
            'debt_to_ebitda': {'healthy': 2.0, 'watch': 4.0, 'inverse': True},
            'current_ratio': {'healthy': 2.0, 'watch': 1.0},
            'quick_ratio': {'healthy': 1.0, 'watch': 0.5},
            'net_income_margin': {'healthy': 10.0, 'watch': 5.0},
            'interest_coverage': {'healthy': 3.0, 'watch': 2.0},
            'leverage_ratio': {'healthy': 0.3, 'watch': 0.6, 'inverse': True},
            'working_capital': {'healthy': 0, 'watch': 0},
            'dso': {'healthy': 45, 'watch': 60, 'inverse': True}
        }

        if ratio_key not in thresholds:
            return 'unknown'

        threshold = thresholds[ratio_key]
        is_inverse = threshold.get('inverse', False)

        if is_inverse:
            # For ratios where lower is better
            if value <= threshold['healthy']:
                return 'healthy'
            elif value <= threshold.get('watch', float('inf')):
                return 'watch'
            else:
                return 'concern'
        else:
            # For ratios where higher is better
            if value >= threshold['healthy']:
                return 'healthy'
            elif value >= threshold.get('watch', 0):
                return 'watch'
            else:
                return 'concern'


# Standalone function for easy integration
def generate_credit_memo_docx(extracted_data, ratios, memo_narrative,
                              borrower_info, output_filename):
    """
    Generate credit memo Word document

    Args:
        extracted_data: Financial data from ADE
        ratios: Calculated financial ratios
        memo_narrative: Text narrative (optional, for reference)
        borrower_info: Borrower details
        output_filename: Output file path

    Returns:
        Path to generated .docx file
    """
    generator = CreditMemoWordGenerator()
    return generator.generate_credit_memo(
        extracted_data=extracted_data,
        ratios=ratios,
        memo_narrative=memo_narrative,
        borrower_info=borrower_info,
        output_path=output_filename
    )


# Test function
if __name__ == "__main__":
    # Test data
    test_extracted = {
        'revenue': 5000000,
        'net_income': 400000,
        'ebitda': 600000,
        'total_assets': 3000000,
        'total_liabilities': 1200000,
        'current_assets': 1500000,
        'current_liabilities': 500000,
        'total_debt': 1000000,
        'interest_expense': 50000
    }

    test_ratios = {
        'dscr': 1.71,
        'debt_to_ebitda': 1.67,
        'current_ratio': 3.00,
        'quick_ratio': 2.00,
        'net_income_margin': 8.0,
        'interest_coverage': 12.00,
        'leverage_ratio': 0.33,
        'working_capital': 1000000,
        'dso': 73,
        'interpretations': {
            'dscr': 'Strong - Excellent debt coverage',
            'debt_to_ebitda': 'Strong - Low leverage',
            'current_ratio': 'Strong liquidity position',
            'quick_ratio': 'Good immediate liquidity',
            'net_income_margin': 'Moderate profitability',
            'interest_coverage': 'Strong - Excellent coverage',
            'leverage_ratio': 'Low leverage - Conservative',
            'working_capital': 'Positive - Good operational liquidity',
            'dso': 'Slow collections - Risk concern'
        }
    }

    test_borrower = {
        'borrower_name': 'Tech Innovations LLC',
        'industry': 'Software Development',
        'loan_amount': 750000,
        'loan_type': 'Commercial Term Loan',
        'purpose': 'Equipment acquisition and working capital',
        'loan_officer': 'Jennifer Martinez',
        'credit_analyst': 'David Chen'
    }

    output_path = "test_credit_memo.docx"

    print("Generating test credit memo...")
    result = generate_credit_memo_docx(
        extracted_data=test_extracted,
        ratios=test_ratios,
        memo_narrative="",
        borrower_info=test_borrower,
        output_filename=output_path
    )

    print(f"✓ Credit memo generated: {result}")
    print("Open the file in Microsoft Word to review.")
